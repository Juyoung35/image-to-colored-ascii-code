import argparse
from PIL import Image
import math
from functools import reduce
from docx import Document
from docx.text.run import Font
from docx.dml.color import ColorFormat
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

MAX_COLOR_VALUE = 255
DOCX_LINE_SPACING_RATIO = 0.55 # 473 / 863

def get_resized_size(width_and_height: tuple, scale):
    return tuple(map(lambda x: int(x * scale), width_and_height))

# resize image if scale option is given
def resize_image(image, scale=1.0, resample='LANCZOS'):
	if scale == 1.0: return image
	resized_image = image.resize(get_resized_size(image.size, scale), resample=resample)
	return resized_image

# convert image to grayscale
def grayscale_image(image):
	grayscaled_image = image.convert(mode="L")
	return grayscaled_image

# write non-colored raw text
def write_raw_txt(dest, image, ascii_chars):
	width, height = image.size
	pixels = image.getdata()
	divider = math.ceil(MAX_COLOR_VALUE / (len(ascii_chars) - 1))
	raw_txt = "".join([ascii_chars[pixel // divider] for pixel in pixels])
	with open(f"{dest}.txt", "w") as f:
		for y in range(height):
			z = y * width
			f.write(raw_txt[z:z + width])
			f.write('\n')

# write colored rich text format
def write_rtf(dest, image, ascii_chars, palette, font_size):
	rtf_txt = ""
	rtf_txt += r"{\rtf1\ansi\ansicpg949\deff0\nouicompat\deflangfe1042{\fonttbl{\f0\fnil\fcharset0 Consolas;}}" + '\n'
	rtf_txt += r"{\colortbl ;"
	for r in range(palette):
		for g in range(palette):
			for b in range(palette):
				red, green, blue = tuple(map(lambda c: MAX_COLOR_VALUE * c // palette, (r, g, b)))
				rtf_txt += f"\\red{red}\\green{green}\\blue{blue};"
	rtf_txt += '}' + '\n'
	rtf_txt += r"{\*\generator Riched20 10.0.19041}\viewkind4\uc1" + '\n'
	rtf_txt += f"\\pard\\sa200\\sl276\\slmult1\\f0\\fs{font_size}\\lang18\\b"

	width, height = image.size
	grayscaled_image = grayscale_image(image)
	pixels = image.getdata()
	grayscale_pixels = grayscaled_image.getdata()
	divider = math.ceil(MAX_COLOR_VALUE / (len(ascii_chars) - 1))
	for y in range(height):
		for x in range(width):
			z = y * width + x
			grayscale_pixel = grayscale_pixels[z]
			pixel = pixels[z][:3]
			r, g, b = pixel
			cf = reduce(lambda acc, t: acc + int(t[1] / (MAX_COLOR_VALUE / palette) * palette**t[0]), enumerate((b, g, r)), 1)
			char = ascii_chars[grayscale_pixel // divider]
			rtf_txt += f"\\cf{cf} {char}"
		rtf_txt += r"\line "

	rtf_txt += '}'
	with open(f"{dest}.rtf", "w") as f:
		f.write(rtf_txt)

def write_docx(dest, image, ascii_chars, font_size, line_spacing):
	document = Document()
	styles = document.styles['Normal']
	font = styles.font
	font.name = 'Consolas'
	font.size = Pt(font_size)
	p = document.add_paragraph()
	p_format = p.paragraph_format
	p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p_format.left_indent = Pt(0)
	p_format.right_indent = Pt(0)
	p_format.space_before = Pt(0)
	p_format.space_after = Pt(0)
	p_format.line_spacing = Pt(font_size * (1 if line_spacing else DOCX_LINE_SPACING_RATIO))

	width, height = image.size
	grayscaled_image = grayscale_image(image)
	pixels = image.getdata()
	grayscale_pixels = grayscaled_image.getdata()
	divider = math.ceil(MAX_COLOR_VALUE / (len(ascii_chars) - 1))
	for y in range(height):
		for x in range(width):
			z = y * width + x
			grayscale_pixel = grayscale_pixels[z]
			pixel = pixels[z][:3]
			r, g, b = pixel
			char = ascii_chars[grayscale_pixel // divider]
			run = p.add_run(char)
			font = run.font
			font.color.rgb = RGBColor(r, g, b)
		p.add_run('\n')

	document.save(f"{dest}.docx")

def write_html(dest, image, ascii_chars, font_size, line_spacing):
	html_txt = ""
	html_txt += "<!DOCTYPE HTML>" + '\n'
	html_txt += "<html>" + '\n'
	html_txt += "<body>" + '\n'
	html_txt += "<style type=\"text/css\">" + '\n'
	html_txt += "\tspan {" + '\n'
	html_txt += f"\t\tfont-size: {font_size}px;" + '\n'
	html_txt += "\t\tfont-weight: bold;" + '\n'
	html_txt += "\t\tfont-family: monospace;" + '\n'
	html_txt += '\t}' + '\n'
	html_txt += "</style>" + '\n'
	
	width, height = image.size
	grayscaled_image = grayscale_image(image)
	pixels = image.getdata()
	grayscale_pixels = grayscaled_image.getdata()
	divider = math.ceil(MAX_COLOR_VALUE / (len(ascii_chars) - 1))
	for y in range(height):
		for x in range(width):
			z = y * width + x
			grayscale_pixel = grayscale_pixels[z]
			pixel = pixels[z][:3]
			r, g, b = pixel
			char = ascii_chars[grayscale_pixel // divider]
			html_txt += f"<span style=\"color: #{r:02x}{g:02x}{b:02x}\">{char}</span>"
		html_txt += "<br>"

	html_txt += "</body>" + '\n'
	html_txt += "</html>"
	with open(f"{dest}.html", "w") as f:
		f.write(html_txt)

def main():
	ascii_chars_file = open("ascii_characters.txt", "r")
	ascii_chars = ascii_chars_file.readlines()
	for i in range(len(ascii_chars)):
		ascii_chars[i] = ascii_chars[i][0]

	parser = argparse.ArgumentParser()
	parser.add_argument(nargs=1, dest="SOURCE", help="the path and name of source file", default="input/sample-image-0.jpg", type=str)
	parser.add_argument("-s", "--scale", nargs='?', dest="SCALE", help="resize SOURCE image file by SCALE. if not given, size will not be changed", default=1.0, type=float)
	parser.add_argument("-d", "--dst", "--dest", "--destination", nargs='?', dest="DEST", help="set output file name with DEST.", default=None, type=str)
	parser.add_argument("-f", "--format", nargs='?', dest="FORMAT", help="get output file format. if not given, format will not be changed.", default=None, type=str)
	parser.add_argument("-p", "--palette", nargs='?', dest="PALETTE", help="limit the palette color in PALETTE^3 numbers if output file format support color", default=8, type=int)
	parser.add_argument("-rs", "--resample", nargs='?', dest="RESAMPLE", help="resampleing filter used in resizing the image. Available filters: LANCZOS(default), BICUBIC, HAMMING, BILINEAR, BOX, NEAREST", default='LANCZOS', type=str)
	parser.add_argument("-fs", "--fontsize", nargs='?', dest="FONTSIZE", help="set size of the font.", default=10, type=int)
	parser.add_argument("-ls", "--linespacing", dest="LINESPACING", help="line spacing excatly as font size in .docx file format. it will size up the height, not the width. if not, image ratio will be same.", default=False, action='store_true')

	src = parser.parse_args().SOURCE[0]
	try:
		image = Image.open(src)
	except:
		print(f"{src} is not a valid image source file name.")
	name = src.split('.')[:-1][0]
	ext = src.split('.')[-1]

	scale = parser.parse_args().SCALE
	dest = parser.parse_args().DEST
	format = parser.parse_args().FORMAT
	if not dest: dest = name if ext != format else f"{name}-output"
	if not format: format = ext
	palette = parser.parse_args().PALETTE
	resample = parser.parse_args().RESAMPLE
	if False: pass
	elif resample == "NEAREST": resample = Image.NEAREST
	elif resample == "LANCZOS": resample = Image.LANCZOS
	elif resample == "BILINEAR": resample = Image.BILINEAR
	elif resample == "BICUBIC": resample = Image.BICUBIC
	elif resample == "BOX": resample = Image.BOX
	elif resample == "HAMMING": resample = Image.HAMMING
	font_size = parser.parse_args().FONTSIZE
	line_spacing = parser.parse_args().LINESPACING

	if False: pass
	elif format == "txt":
		write_raw_txt(dest, grayscale_image(resize_image(image, scale, resample)), ascii_chars)
	elif format == "rtf":
		write_rtf(dest, resize_image(image, scale, resample), ascii_chars, palette, font_size << 1)
	elif format == "docx":
		write_docx(dest, resize_image(image, scale, resample), ascii_chars, font_size, line_spacing)
	elif format == "html":
		write_html(dest, resize_image(image, scale, resample), ascii_chars, font_size, line_spacing)
	ascii_chars_file.close()

if __name__ == '__main__':
	main()